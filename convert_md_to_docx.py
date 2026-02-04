import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_doc(md_file, docx_file):
    doc = Document()
    
    # Set style for the whole document or specific parts if needed
    # For now, we'll just code the logic to parse MD lines
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        
        if not line or line == '---':
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Checklists / Bullets
        elif line.startswith('- [ ]') or line.startswith('- [x]'):
            # Checklist item
            status = "✅ " if line.startswith('- [x]') else "⬜ "
            text = line[6:]
            # Bold handling for **Text**
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, status + text)
            
        elif line.startswith('- '):
            # Normal bullet
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            _add_formatted_text(p, text)
            
        # Bold text lines (like Goal: ... Status: ...)
        elif line.startswith('**Goal:**') or line.startswith('**Status:**'):
             p = doc.add_paragraph()
             _add_formatted_text(p, line)

        else:
            # Normal paragraph
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            
    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

def _add_formatted_text(paragraph, text):
    # Split by bold markers **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    create_word_doc("PROJECT_PHASES.md", "PROJECT_PHASES.docx")
